"""
Handles text extraction from uploaded documents.

Native-text PDFs/DOCX are parsed directly. Scanned documents (or PDFs with
no extractable text layer) fall back to OCR.
"""
import io

from pypdf import PdfReader
from docx import Document as DocxDocument
from pdf2image import convert_from_bytes
import pytesseract
from PIL import Image


def extract_text(raw_bytes: bytes, suffix: str) -> tuple[str, bool]:
    """
    Extracts text from a document's raw bytes.

    Returns:
        (text, ocr_used) — the extracted text, and whether OCR was required.
    """
    if suffix == ".docx":
        return _extract_docx(raw_bytes), False

    if suffix == ".pdf":
        text = _extract_pdf_native(raw_bytes)
        if _looks_like_scanned(text):
            return _extract_pdf_ocr(raw_bytes), True
        return text, False

    if suffix in {".png", ".jpg", ".jpeg"}:
        return _extract_image_ocr(raw_bytes), True

    if suffix == ".txt":
        return raw_bytes.decode("utf-8", errors="ignore"), False

    raise ValueError(f"Unsupported file type: {suffix}")


def _extract_docx(raw_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(raw_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_pdf_native(raw_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(raw_bytes))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        pages.append(f"\n[[PAGE {i}]]\n{page_text}")
    return "".join(pages)


def _looks_like_scanned(text: str) -> bool:
    return len(text.strip()) < 50


def _extract_pdf_ocr(raw_bytes: bytes) -> str:
    images = convert_from_bytes(raw_bytes, dpi=200)
    pages = []
    for i, image in enumerate(images, start=1):
        page_text = pytesseract.image_to_string(image)
        pages.append(f"\n[[PAGE {i}]]\n{page_text}")
    return "".join(pages)


def _extract_image_ocr(raw_bytes: bytes) -> str:
    image = Image.open(io.BytesIO(raw_bytes))
    return f"\n[[PAGE 1]]\n{pytesseract.image_to_string(image)}"
